"""
Word 导入多模板解析器。

支持两种模板格式：
- 模板A（标准检查表）：表格逐行列出隐患，每行一个。
- 模板B（通知单）：按"一、"、"二、"分项目，按"1."、"2."分隐患，图片嵌入段落。

自动检测模板类型并调用对应解析逻辑。
"""

import re
from docx import Document
from lxml import etree
from typing import List, Tuple, Optional
import io
from datetime import date as _date, datetime as _datetime, timedelta as _timedelta


def parse_deadline(deadline_str: str, base_date=None):
    """
    解析期限字符串，支持格式：
    - 2026-06-01
    - 2026.6.1
    - 2026/6/1
    - 2026年6月1日
    - 4月3日 / 4月7日前（年份从检查日期推断）
    - 立即整改（按检查当天）
    - 3日内 / 3天内 / 3个工作日内（按检查日期顺延自然日）
    """
    if not deadline_str:
        return None

    text = str(deadline_str).strip()
    text = text.replace("．", ".").replace("－", "-").replace("／", "/")
    text = re.sub(r"\s+", "", text)

    if isinstance(base_date, str):
        base_date = parse_check_date(base_date)
    if not isinstance(base_date, _date):
        base_date = None

    if "立即整改" in text or "即刻整改" in text:
        return base_date or _date.today()

    patterns = [
        r'(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})',
        r'(\d{4})年(\d{1,2})月(\d{1,2})日',
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            try:
                y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
                return _date(y, mo, d)
            except ValueError:
                continue

    m = re.search(r'(\d{1,2})月(\d{1,2})日?(?:前|之前|以前)?', text)
    if m:
        try:
            year = base_date.year if base_date else _date.today().year
            candidate = _date(year, int(m.group(1)), int(m.group(2)))
            if base_date and candidate < base_date:
                candidate = _date(year + 1, int(m.group(1)), int(m.group(2)))
            return candidate
        except ValueError:
            pass

    m = re.search(r'(\d+)(?:个)?(?:工作日|天|日)内', text)
    if m:
        start = base_date or _date.today()
        return start + _timedelta(days=int(m.group(1)))

    return None


def parse_check_date(text: str):
    """
    解析检查日期，支持：
    - 检查时间：2026.3.31
    - 检查日期：2026-03-31
    - 2026年2月26日
    """
    if not text:
        return None

    patterns = [
        r'检查(?:时间|日期)\s*[：:]\s*(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})',
        r'检查(?:时间|日期)\s*[：:]\s*(\d{4})年\s*(\d{1,2})月\s*(\d{1,2})日',
        r'(\d{4})年\s*(\d{1,2})月\s*(\d{1,2})日',
        r'(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if not match:
            continue
        try:
            return _date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
        except ValueError:
            continue
    return None


# ── 模板B（通知单格式）解析 ────────────────────────────────────────────────────

def extract_images_from_cell(cell, issue_id: int) -> list:
    """
    从表格单元格中提取所有图片。
    返回：[(file_name, image_data), ...]
    """
    images = []
    for para_idx, para in enumerate(cell.paragraphs):
        drawing = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        if drawing is not None:
            for blip in drawing.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                rel_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rel_id:
                    try:
                        image_data = para._parent.part.rels[rel_id].target_part.blob
                        file_name = f"cell_img_{issue_id}_{para_idx}.jpg"
                        images.append((file_name, image_data))
                    except Exception:
                        pass
    return images


def extract_images_from_docx(doc) -> dict:
    """
    从 docx 中提取所有图片的二进制数据。
    返回：{rel_id: image_bytes}
    """
    images = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            try:
                image_data = rel.target_part.blob
                images[rel_id] = image_data
            except Exception:
                pass
    return images


def parse_template_b(doc) -> List[Tuple[str, List[dict]]]:
    """
    解析通知单格式。

    结构：
    - 表格1：表头信息（项目名称、责任部门等）
    - 行3（"主要问题"行）的合并单元格（B:D列）包含全文：
      "一、项目A"
         "1. 问题xxx"
         [图片]
         "2. 问题xxx"
      "二、项目B"
         "1. 问题xxx"

    返回：[(project_name, [{"title": ..., "description": ..., "image_rel_ids": ...}])]
    """
    if not doc.tables:
        return []

    # 1. 从表格1读取所有项目名称（B1单元格，用"、"分隔）
    table = doc.tables[0]
    try:
        project_cell_text = table.cell(0, 1).text.strip()  # B1单元格
        project_names = [p.strip() for p in project_cell_text.split("、") if p.strip()]
    except Exception:
        project_names = ["未分类"]

    # 2. 找到"主要问题"行的合并单元格（只取第一个匹配的单元格）
    content_cell = None
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            if row_idx != 3:  # 只要第4行（索引3）
                continue
            # 行3的合并单元格（B:D列）→ 取B列（索引1）
            content_cell = row.cells[1]
            break
        if content_cell:
            break

    if not content_cell:
        return []

    # 3. 抓检查日期（从开头介绍文字或表格字段里抓 "YYYY年MM月DD日"）
    check_date = None
    for para in content_cell.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        check_date = parse_check_date(text)
        if check_date:
            break
    if not check_date:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            for idx, cell_text in enumerate(cells):
                if "检查时间" in cell_text or "检查日期" in cell_text:
                    right_text = cells[idx + 1] if idx + 1 < len(cells) else ""
                    check_date = parse_check_date(right_text or cell_text)
                    if check_date:
                        break
            if check_date:
                break
    check_date_str = check_date.isoformat() if check_date else None

    # 4. 模板B通常是整张通知单一个统一整改期限
    deadline = None
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        for idx, cell_text in enumerate(cells):
            if "整改期限" in re.sub(r"\s+", "", cell_text):
                right_text = cells[idx + 1] if idx + 1 < len(cells) else ""
                deadline = parse_deadline(right_text or cell_text, check_date)
                break
        if deadline:
            break
    deadline_str = deadline.isoformat() if deadline else None

    # 5. 解析单元格中的所有段落
    # 项目名直接从"一、项目A"等文本中抓取，不用 B1 单元格
    current_project = ""
    current_issues = []
    current_issue = None
    results = []

    for para in content_cell.paragraphs:
        text = para.text.strip()
        has_img = para._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') is not None

        # 1. 先处理图片（attach 到当前问题）
        # 重要：即使段落里有"二、项目名"文字，图片也应该属于前一个问题
        if has_img and current_issue:
            drawings = para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
            for drawing in drawings:
                for blip in drawing.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                    rel_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rel_id:
                        if 'image_rel_ids' not in current_issue:
                            current_issue['image_rel_ids'] = []
                        current_issue['image_rel_ids'].append(rel_id)

        # 2. 再处理文字
        if text:
            project_match = re.search(r'^([一二三四五六七八九十]+)[、.]\s*(.+)$', text)
            if project_match:
                # 项目切换：先保存当前问题
                if current_issue:
                    current_issues.append(current_issue)
                # 再保存当前项目
                if current_issues:
                    results.append((current_project, current_issues))
                # 新项目名 = "一、项目A" → "项目A"
                current_project = project_match.group(2).strip()
                current_issues = []
                current_issue = None
            else:
                # 检查是否是新问题（"1."、"1、" 格式，必须有点或顿号分隔）
                issue_match = re.search(r'^(\d+)[.、．]\s*(.+)', text)
                if issue_match:
                    # 保存上一个问题
                    if current_issue:
                        current_issues.append(current_issue)
                    # 开始新问题
                    current_issue = {
                        "title": issue_match.group(2).strip(),
                        "description": "",
                        "notes": "",
                        "deadline": deadline_str,
                        "image_rel_ids": [],
                        "check_date": check_date_str
                    }
                elif current_issue:
                    # 普通文本追加到描述
                    if current_issue["description"]:
                        current_issue["description"] += "\n" + text
                    else:
                        current_issue["description"] = text

    # 保存最后一个问题
    if current_issue:
        current_issues.append(current_issue)

    # 保存最后一个项目
    if current_issues:
        results.append((current_project, current_issues))

    return results


def _iter_body_elements(doc):
    """按文档XML顺序迭代 body 的子元素，返回 (tag, element) 对。"""
    body = doc.element.body
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        yield tag, child


def parse_template_a(doc) -> List[Tuple[str, List[dict]]]:
    """
    解析标准检查表格式（表格逐行列出隐患）。

    按文档XML顺序遍历 body：
    - 遇到段落（<w:p>）：如果含"企业名称"/"项目名称"，更新 current_project
    - 遇到表格（<w:tbl>）：解析该表格，归到 current_project

    支持一个Word文件含多个项目（每个项目前有"企业名称：XXX"段落）。
    返回：[(project_name, [{"title": ..., "check_date": ...}])]
    """
    current_project = ""
    current_check_date = None
    table_cursor = 0  # doc.tables 的索引
    results = []  # [(project_name, items)]
    current_items = []

    for tag, element in _iter_body_elements(doc):
        if tag == 'p':  # 段落
            # 拼接段落全文（可能含换行）
            texts = [node.text for node in element.iter() if node.text and node.text.strip()]
            text = ' '.join(texts)
            if not text:
                continue

            parsed_check_date = parse_check_date(text)
            if parsed_check_date:
                current_check_date = parsed_check_date

            # 检查是否是新项目
            for prefix in ["企业名称", "项目名称"]:
                if prefix in text:
                    val = re.search(
                        r'' + prefix + r'[：:]\s*(.+?)(?:\s{2,}隐患|\s{2,}检查|\s*检查(?:时间|日期)|$)',
                        text, re.DOTALL
                    )
                    if val:
                        # 保存上一个项目
                        if current_items:
                            results.append((current_project, current_items))
                        current_project = val.group(1).strip()
                        current_items = []
                        break

        elif tag == 'tbl':  # 表格
            if table_cursor >= len(doc.tables):
                break
            table = doc.tables[table_cursor]
            table_cursor += 1

            for row_idx, row in enumerate(table.rows):
                if row_idx == 0:
                    continue
                cells = row.cells
                if len(cells) < 3:
                    continue
                title = cells[2].text.strip() if len(cells) >= 3 else ""
                description = cells[3].text.strip() if len(cells) >= 4 else ""
                notes = cells[4].text.strip() if len(cells) >= 5 else ""
                remarks = cells[5].text.strip() if len(cells) >= 6 else ""
                deadline = parse_deadline(remarks, current_check_date) if remarks else None
                if not title:
                    continue
                current_items.append({
                    "title": title[:200],
                    "description": description[:500] if description else None,
                    "notes": notes[:500] if notes else None,
                    "deadline": deadline.isoformat() if deadline else None,
                    "check_date": current_check_date.isoformat() if current_check_date else None,
                    "_table_idx": table_cursor - 1,  # 当前表格索引（0-based）
                    "_row_idx": row_idx,              # 当前行索引
                })

    # 保存最后一个项目
    if current_items:
        results.append((current_project, current_items))

    return results

def detect_template(doc) -> str:
    """
    自动检测 Word 文档使用的模板类型。
    """
    if not doc.tables:
        return "unknown"

    # 检查是否有"主要问题"标签（模板B特征）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text_normalized = re.sub(r'\s+', '', cell.text)
                if "主要问题" in cell_text_normalized:
                    # 检查同一行其他单元格是否含"一、"格式
                    for content_cell in row.cells:
                        for para in content_cell.paragraphs:
                            if re.search(r'^[一二三四五六七八九十]+[、.]', para.text.strip()):
                                return "template_b"

    # 检查是否是标准检查表格式（模板A特征）
    for table in doc.tables:
        if table.rows:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            if any(h in headers for h in ["序号", "问题描述", "整改建议", "检查项目"]):
                return "template_a"

    return "template_a"


def parse_word_doc(doc, template_type: Optional[str] = None) -> List[Tuple[str, List[dict]]]:
    """
    统一入口：自动检测模板类型并调用对应解析器。
    """
    if template_type is None:
        template_type = detect_template(doc)

    if template_type == "template_b":
        return parse_template_b(doc)
    elif template_type == "template_a":
        return parse_template_a(doc)
    else:
        # 默认尝试模板A
        return parse_template_a(doc)
