from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from typing import List, Optional
from pydantic import BaseModel
from datetime import date, datetime
import os
import uuid
from PIL import Image
import io
import re
import hashlib

from app.database import get_db
from app.models import SafetyIssue, Photo
from app.schemas import IssueCreate, IssueUpdate, IssueResponse
from app.utils import save_upload_file
from app.api.word_parser import (
    parse_word_doc,
    extract_images_from_docx,
    extract_images_from_cell as extract_images_from_docx_cell,
)  # 多模板解析器

router = APIRouter(prefix="/api/issues", tags=["问题管理"])

APP_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
PHOTOS_ROOT = os.path.join(APP_ROOT, "data", "photos")
VALID_STATUSES = {"待整改", "整改中", "已完成"}
COMPLETED_STATUSES = {"已完成", "已整改"}


class BatchStatusUpdate(BaseModel):
    issue_ids: List[int]
    status: str


def _parse_check_date_value(value):
    if not value:
        return None
    if isinstance(value, date):
        return value
    if isinstance(value, str):
        try:
            return date.fromisoformat(value[:10])
        except ValueError:
            return None
    return None


def _is_overdue(issue: SafetyIssue) -> bool:
    return bool(
        issue.deadline
        and issue.deadline < date.today()
        and (issue.status or "") not in COMPLETED_STATUSES
    )


def _has_rectification_photo(issue: SafetyIssue) -> bool:
    return any(photo.photo_type == "整改照片" for photo in issue.photos)


def _filtered_issue_query(
    db: Session,
    status: Optional[str] = None,
    severity: Optional[str] = None,
    project_name: Optional[str] = None,
    check_date: Optional[date] = None,
    overdue: Optional[bool] = None,
):
    query = db.query(SafetyIssue)

    if status:
        query = query.filter(SafetyIssue.status == status)
    if severity:
        query = query.filter(SafetyIssue.severity == severity)
    if project_name:
        query = query.filter(SafetyIssue.project_name == project_name)
    if check_date:
        query = query.filter(SafetyIssue.check_date == check_date)
    if overdue is True:
        query = query.filter(
            SafetyIssue.deadline.isnot(None),
            SafetyIssue.deadline < date.today(),
            SafetyIssue.status.notin_(list(COMPLETED_STATUSES)),
        )

    return query


def _build_import_key(index: int, project_name: str, check_date_value, title: str) -> str:
    check_date_str = check_date_value.isoformat() if isinstance(check_date_value, date) else str(check_date_value or "")
    raw = f"{index}|{project_name or ''}|{check_date_str}|{(title or '').strip()}"
    digest = hashlib.sha1(raw.encode("utf-8")).hexdigest()[:16]
    return f"{index}_{digest}"


def _duplicate_issue_count(db: Session, project_name: str, check_date_value, title: str) -> int:
    clean_title = (title or "").strip()[:200]
    if not clean_title:
        return 0

    query = db.query(SafetyIssue).filter(SafetyIssue.title == clean_title)
    if project_name:
        query = query.filter(SafetyIssue.project_name == project_name)
    else:
        from sqlalchemy import or_
        query = query.filter(or_(SafetyIssue.project_name.is_(None), SafetyIssue.project_name == ""))

    parsed_date = _parse_check_date_value(check_date_value)
    if parsed_date:
        query = query.filter(SafetyIssue.check_date == parsed_date)
    else:
        query = query.filter(SafetyIssue.check_date.is_(None))

    return query.count()


def save_photo_bytes(issue_id: int, file_name: str, image_data: bytes) -> tuple[str, str]:
    """保存从 Word 中提取的图片，返回绝对路径和入库文件名。"""
    issue_dir = os.path.join(PHOTOS_ROOT, str(issue_id))
    os.makedirs(issue_dir, exist_ok=True)

    safe_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", file_name).strip("._")
    if not safe_name:
        safe_name = f"{uuid.uuid4().hex}.jpg"
    stored_name = f"{uuid.uuid4().hex[:8]}_{safe_name}"
    file_path = os.path.join(issue_dir, stored_name)

    with open(file_path, "wb") as f:
        f.write(image_data)
    return file_path, stored_name

@router.get("/")
def get_issues(
    page: int = 1,
    limit: int = 100,
    status: Optional[str] = None,
    severity: Optional[str] = None,
    project_name: Optional[str] = None,
    check_date: Optional[date] = None,
    overdue: Optional[bool] = None,
    db: Session = Depends(get_db)
):
    query = _filtered_issue_query(db, status, severity, project_name, check_date, overdue)

    total = query.count()
    skip = (page - 1) * limit
    issues = query.order_by(SafetyIssue.id.asc()).offset(skip).limit(limit).all()
    return {"items": [issue.to_dict() for issue in issues], "total": total}


@router.get("/projects/list")
def get_projects(db: Session = Depends(get_db)):
    """获取所有不重复的项目名称"""
    rows = db.query(SafetyIssue.project_name).filter(
        SafetyIssue.project_name.isnot(None),
        SafetyIssue.project_name != ""
    ).distinct().all()
    return [r[0] for r in rows]


@router.get("/stats/summary")
def get_issue_stats(
    status: Optional[str] = None,
    severity: Optional[str] = None,
    project_name: Optional[str] = None,
    check_date: Optional[date] = None,
    overdue: Optional[bool] = None,
    db: Session = Depends(get_db)
):
    """按当前筛选条件统计隐患、逾期和整改照片完整度。"""
    issues = _filtered_issue_query(db, status, severity, project_name, check_date, overdue).all()
    total = len(issues)
    overdue_count = sum(1 for issue in issues if _is_overdue(issue))
    completed_count = sum(1 for issue in issues if (issue.status or "") in COMPLETED_STATUSES)
    rectification_photo_count = sum(1 for issue in issues if _has_rectification_photo(issue))

    return {
        "total": total,
        "overdue_count": overdue_count,
        "pending_count": total - completed_count,
        "completed_count": completed_count,
        "rectification_photo_count": rectification_photo_count,
        "rectification_missing_count": total - rectification_photo_count,
    }

@router.get("/{issue_id}", response_model=IssueResponse)
def get_issue(issue_id: int, db: Session = Depends(get_db)):
    issue = db.query(SafetyIssue).filter(SafetyIssue.id == issue_id).first()
    if not issue:
        raise HTTPException(status_code=404, detail="问题不存在")
    return issue.to_dict()

@router.post("/", response_model=IssueResponse)
def create_issue(issue: IssueCreate, db: Session = Depends(get_db)):
    db_issue = SafetyIssue(**issue.dict())
    db.add(db_issue)
    db.commit()
    db.refresh(db_issue)
    return db_issue.to_dict()

@router.post("/with-photos", response_model=IssueResponse)
async def create_issue_with_photos(
    title: str = Form(...),
    description: Optional[str] = Form(None),
    location: Optional[str] = Form(None),
    severity: str = Form("一般"),
    responsible_person: Optional[str] = Form(None),
    deadline: Optional[str] = Form(None),
    notes: Optional[str] = Form(None),
    photos: List[UploadFile] = File([]),
    db: Session = Depends(get_db)
):
    db_issue = SafetyIssue(
        title=title,
        description=description,
        location=location,
        severity=severity,
        responsible_person=responsible_person,
        deadline=date.fromisoformat(deadline) if deadline else None,
        notes=notes
    )
    db.add(db_issue)
    db.commit()
    db.refresh(db_issue)

    for photo in photos:
        if photo.filename:
            file_path, file_name = await save_upload_file(photo, db_issue.id, "问题照片")
            db_photo = Photo(
                issue_id=db_issue.id,
                photo_type="问题照片",
                file_path=file_path,
                file_name=file_name
            )
            db.add(db_photo)

    db.commit()
    db.refresh(db_issue)
    return db_issue.to_dict()

@router.put("/{issue_id}", response_model=IssueResponse)
def update_issue(issue_id: int, issue: IssueUpdate, db: Session = Depends(get_db)):
    db_issue = db.query(SafetyIssue).filter(SafetyIssue.id == issue_id).first()
    if not db_issue:
        raise HTTPException(status_code=404, detail="问题不存在")

    update_data = issue.dict(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_issue, key, value)

    db.commit()
    db.refresh(db_issue)
    return db_issue.to_dict()

@router.delete("/{issue_id}")
def delete_issue(issue_id: int, db: Session = Depends(get_db)):
    db_issue = db.query(SafetyIssue).filter(SafetyIssue.id == issue_id).first()
    if not db_issue:
        raise HTTPException(status_code=404, detail="问题不存在")

    _delete_issue_photo_files(db_issue)

    db.delete(db_issue)
    db.commit()
    return {"message": "删除成功"}


def _delete_issue_photo_files(issue: SafetyIssue):
    for photo in issue.photos:
        if os.path.exists(photo.file_path):
            os.remove(photo.file_path)


@router.delete("/all/delete")
def delete_all_issues(confirm: str, db: Session = Depends(get_db)):
    """删除全部项目和隐患。confirm 必须为 DELETE_ALL，防止误删。"""
    if confirm != "DELETE_ALL":
        raise HTTPException(status_code=400, detail="缺少删除确认参数")

    issues = db.query(SafetyIssue).all()
    if not issues:
        return {"message": "当前没有可删除的问题", "deleted_count": 0}

    deleted_count = 0
    for issue in issues:
        _delete_issue_photo_files(issue)
        db.delete(issue)
        deleted_count += 1

    db.commit()
    return {"message": f"成功删除全部 {deleted_count} 条问题", "deleted_count": deleted_count}


@router.delete("/batch/delete")
def batch_delete_issues(
    project_name: Optional[str] = None,
    check_date: Optional[date] = None,
    issue_ids: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """批量删除：按项目名删除 或 按 ID 列表删除"""
    if not project_name and not check_date and not issue_ids:
        raise HTTPException(status_code=400, detail="请提供删除条件")

    query = db.query(SafetyIssue)
    
    if project_name:
        # "未分类" 特别处理：匹配 NULL/空字符串
        if project_name in ('未分类', '__EMPTY__'):
            from sqlalchemy import or_
            query = query.filter(or_(
                SafetyIssue.project_name.is_(None),
                SafetyIssue.project_name == '',
                SafetyIssue.project_name == '未分类'
            ))
        else:
            query = query.filter(SafetyIssue.project_name == project_name)

    if check_date:
        query = query.filter(SafetyIssue.check_date == check_date)
    
    if issue_ids:
        ids = []
        for sid in issue_ids.split(','):
            try:
                ids.append(int(sid.strip()))
            except ValueError:
                pass
        if ids:
            query = query.filter(SafetyIssue.id.in_(ids))
    
    issues = query.all()
    if not issues:
        raise HTTPException(status_code=404, detail="没有可删除的问题")
    
    deleted_count = 0
    for issue in issues:
        _delete_issue_photo_files(issue)
        db.delete(issue)
        deleted_count += 1
    
    db.commit()
    return {"message": f"成功删除 {deleted_count} 条问题", "deleted_count": deleted_count}


@router.post("/{issue_id}/photos", response_model=IssueResponse)
async def upload_photo(
    issue_id: int,
    photo_type: str = Form(...),
    photos: List[UploadFile] = File(...),
    db: Session = Depends(get_db)
):
    db_issue = db.query(SafetyIssue).filter(SafetyIssue.id == issue_id).first()
    if not db_issue:
        raise HTTPException(status_code=404, detail="问题不存在")

    for photo in photos:
        if photo.filename:
            file_path, file_name = await save_upload_file(photo, issue_id, photo_type)
            db_photo = Photo(
                issue_id=issue_id,
                photo_type=photo_type,
                file_path=file_path,
                file_name=file_name
            )
            db.add(db_photo)

    db.commit()
    db.refresh(db_issue)
    return db_issue.to_dict()

@router.put("/{issue_id}/status")
def update_status(
    issue_id: int,
    status: str,
    db: Session = Depends(get_db)
):
    db_issue = db.query(SafetyIssue).filter(SafetyIssue.id == issue_id).first()
    if not db_issue:
        raise HTTPException(status_code=404, detail="问题不存在")

    if status not in VALID_STATUSES:
        raise HTTPException(status_code=400, detail="状态值无效")

    db_issue.status = status
    db.commit()
    db.refresh(db_issue)
    return db_issue.to_dict()


@router.put("/batch/update-status")
def batch_update_status(
    payload: BatchStatusUpdate,
    db: Session = Depends(get_db)
):
    """批量更新隐患状态。"""
    if payload.status not in VALID_STATUSES:
        raise HTTPException(status_code=400, detail="状态值无效")

    ids = sorted({int(issue_id) for issue_id in payload.issue_ids if issue_id})
    if not ids:
        raise HTTPException(status_code=400, detail="请选择要更新的隐患")

    issues = db.query(SafetyIssue).filter(SafetyIssue.id.in_(ids)).all()
    if not issues:
        raise HTTPException(status_code=404, detail="没有找到要更新的隐患")

    for issue in issues:
        issue.status = payload.status

    db.commit()
    return {"message": f"已更新 {len(issues)} 条隐患状态", "updated_count": len(issues)}

def parse_deadline(deadline_str):
    if not deadline_str:
        return None
    match = re.search(r'(\d+)月(\d+)日', deadline_str)
    if match:
        month = int(match.group(1))
        day = int(match.group(2))
        year = datetime.now().year
        return date(year, month, day)
    return None

def extract_images_from_cell_legacy(cell, issue_id):
    """从表格单元格中提取图片（精确匹配）"""
    images = []
    try:
        from lxml import etree
        cell_xml = cell._element
        # 查找所有图片关系 ID
        for blip in cell_xml.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
            embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if embed and embed in cell.part.rels:
                rel = cell.part.rels[embed]
                image_part = rel.target_part
                image_data = image_part.blob
                file_ext = os.path.splitext(rel.target_ref)[1] or '.png'
                file_name = f"{issue_id}_{len(images) + 1}{file_ext}"
                images.append((file_name, image_data))
    except Exception as e:
        pass
    return images

def extract_images_from_doc(doc, issue_id):
    """旧版：从文档全局提取图片（备用）"""
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            file_ext = os.path.splitext(rel.target_ref)[1] or '.png'
            file_name = f"{issue_id}_{uuid.uuid4().hex[:8]}{file_ext}"
            images.append((file_name, image_data))
    return images

@router.post("/preview-from-word")
async def preview_from_word(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """仅解析 Word，不写库；返回识别到的项目名称和隐患条目预览。"""
    try:
        from docx import Document

        contents = await file.read()
        doc = Document(io.BytesIO(contents))

        result = parse_word_doc(doc)

        # 展平为 items 列表（带 project_name 字段）
        items = []
        duplicate_count = 0
        item_index = 0
        for project_name, issues in result:
            for issue in issues:
                item_index += 1
                issue["project_name"] = project_name
                check_date_value = _parse_check_date_value(issue.get("check_date"))
                deadline_value = _parse_check_date_value(issue.get("deadline"))
                issue["check_date"] = check_date_value.isoformat() if check_date_value else None
                issue["deadline"] = deadline_value.isoformat() if deadline_value else None
                issue["import_key"] = _build_import_key(item_index, project_name, check_date_value, issue.get("title", ""))
                issue["duplicate_count"] = _duplicate_issue_count(db, project_name, check_date_value, issue.get("title", ""))
                issue["is_duplicate"] = issue["duplicate_count"] > 0
                if issue["is_duplicate"]:
                    duplicate_count += 1
                items.append(issue)

        all_projects = [pn for pn, _ in result]
        first_project = all_projects[0] if all_projects else ""

        return {
            "success": True,
            "project_name": first_project,
            "project_list": all_projects,
            "items": items,
            "duplicate_count": duplicate_count,
            "tables_found": len(doc.tables),
        }
    except Exception as e:
        return {"success": False, "message": str(e)}


@router.post("/import-from-word")
async def import_from_word(
    file: UploadFile = File(...),
    skip_indices: Optional[str] = Form(None),
    skip_duplicates: bool = Form(True),
    db: Session = Depends(get_db)
):
    try:
        from docx import Document

        contents = await file.read()
        doc = Document(io.BytesIO(contents))

        result = parse_word_doc(doc)

        # 提取文档中的所有图片（用于模板B）
        all_images = extract_images_from_docx(doc)

        # 解析要跳过的行索引
        skip_set = set()
        if skip_indices:
            for idx_str in skip_indices.split(','):
                try:
                    skip_set.add(int(idx_str.strip()))
                except ValueError:
                    pass

        imported_count = 0
        skipped_duplicate_count = 0
        errors = []
        item_index = 0

        for project_name, issues in result:
            for issue_idx, issue_dict in enumerate(issues):
                item_index += 1

                check_date_val = _parse_check_date_value(issue_dict.get('check_date'))
                deadline_val = _parse_check_date_value(issue_dict.get('deadline'))
                title = issue_dict.get('title', '')
                skip_key = _build_import_key(item_index, project_name, check_date_val, title)
                if skip_key in skip_set:
                    continue
                if skip_duplicates and _duplicate_issue_count(db, project_name, check_date_val, title) > 0:
                    skipped_duplicate_count += 1
                    continue
                try:
                    db_issue = SafetyIssue(
                        project_name=project_name or None,
                        title=title[:200],
                        description=issue_dict.get('description', '')[:500],
                        location='',
                        severity='一般',
                        deadline=deadline_val,
                        check_date=check_date_val,
                        status='待整改',
                        notes=issue_dict.get('notes', '')[:500],
                        responsible_person='',
                    )
                    db.add(db_issue)
                    db.commit()
                    db.refresh(db_issue)

                    # 保存图片
                    saved_photo = False

                    # 模板B：从 image_rel_ids 提取图片
                    if issue_dict.get('image_rel_ids'):
                        for rel_id in issue_dict['image_rel_ids']:
                            if rel_id in all_images:
                                image_data = all_images[rel_id]
                                file_name = f"issue_{issue_idx}_{rel_id}.jpg"
                                file_path, stored_name = save_photo_bytes(db_issue.id, file_name, image_data)
                                db_photo = Photo(
                                    issue_id=db_issue.id,
                                    photo_type="问题照片",
                                    file_path=file_path,
                                    file_name=stored_name
                                )
                                db.add(db_photo)
                                saved_photo = True

                    # 模板A：从表格单元格提取图片
                    if not saved_photo:
                        try:
                            tbl_idx = issue_dict.get('_table_idx', 0)
                            row_idx = issue_dict.get('_row_idx', 0)
                            if tbl_idx < len(doc.tables):
                                table = doc.tables[tbl_idx]
                                if row_idx < len(table.rows):
                                    row = table.rows[row_idx]
                                    cells = row.cells
                                    if len(cells) >= 2:
                                        cell_b = cells[1]  # 现场图片列
                                        images = extract_images_from_docx_cell(cell_b, db_issue.id)
                                        for file_name, image_data in images:
                                            file_path, stored_name = save_photo_bytes(db_issue.id, file_name, image_data)
                                            db_photo = Photo(
                                                issue_id=db_issue.id,
                                                photo_type="问题照片",
                                                file_path=file_path,
                                                file_name=stored_name
                                            )
                                            db.add(db_photo)
                        except Exception as img_e:
                            errors.append(f"图片提取失败: {str(img_e)}")

                    imported_count += 1
                except Exception as e:
                    db.rollback()
                    errors.append(str(e))

        db.commit()

        all_projects = [pn for pn, _ in result]
        return {
            'success': True,
            'imported_count': imported_count,
            'skipped_duplicate_count': skipped_duplicate_count,
            'errors': errors,
            'project_name': all_projects[0] if all_projects else "",
            'project_list': all_projects,
            'tables_found': len(doc.tables),
        }
    except Exception as e:
        return {'success': False, 'message': str(e)}


# _parse_word_doc 已移至 word_parser.py
